import os
import sys
from docx import Document
from pypdf import PdfReader
from API_Fallback import ask_gemini, ask_openai

def configure_api():
    gemini_key = os.environ.get("GEMINI_API_KEY", "").strip()
    if not gemini_key or "YOUR_" in gemini_key:
        print("Note: GEMINI_API_KEY is not configured with a valid key.")
        user_key = input("Please enter your Gemini API Key (or press Enter to skip): ").strip()
        if user_key:
            os.environ["GEMINI_API_KEY"] = user_key

def ask_llm(query: str) -> str:
    response = ask_gemini(query)
    if "Error contacting Gemini API" in response or "API key not valid" in response:
        alt_response = ask_openai(query)
        if "Error contacting OpenAI API" not in alt_response:
            return alt_response
    return response

def extract_text_from_pdf(file_path: str) -> str:
    reader = PdfReader(file_path)
    text = []
    for page in reader.pages:
        extracted = page.extract_text()
        if extracted:
            text.append(extracted)
    return "\n".join(text)

def extract_text_from_docx(file_path: str) -> str:
    doc = Document(file_path)
    text = []
    for paragraph in doc.paragraphs:
        if paragraph.text:
            text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text for cell in row.cells if cell.text]
            if row_text:
                text.append(" | ".join(row_text))
    return "\n".join(text)

def main():
    configure_api()
    
    if len(sys.argv) > 1:
        file_path = sys.argv[1]
    else:
        file_path = input("\nEnter the path to the PDF or DOCX file: ").strip()
        
    if not os.path.exists(file_path):
        print(f"Error: File not found at '{file_path}'")
        return

    ext = os.path.splitext(file_path)[1].lower()
    print("Scanning document...")
    try:
        if ext == ".pdf":
            doc_text = extract_text_from_pdf(file_path)
        elif ext == ".docx":
            doc_text = extract_text_from_docx(file_path)
        else:
            print("Unsupported file format. Please provide a .pdf or .docx file.")
            return
    except Exception as e:
        print(f"Error reading document: {e}")
        return

    if not doc_text.strip():
        print("Warning: No text could be extracted from the document.")
        return

    print("Document successfully loaded.")
    
    while True:
        try:
            query = input("\nAsk a question (type 'exit' to quit): ").strip()
            if query.lower() in ['exit', 'quit', 'bye']:
                break
            if not query:
                continue

            prompt = (
                f"Based on the following document context, please answer the question.\n\n"
                f"Document Context:\n{doc_text}\n\n"
                f"Question: {query}\n\n"
                f"Answer:"
            )
            print("Generating answer...")
            answer = ask_llm(prompt)
            print(f"\nBot: {answer}")
        except (KeyboardInterrupt, EOFError):
            break

if __name__ == "__main__":
    main()
